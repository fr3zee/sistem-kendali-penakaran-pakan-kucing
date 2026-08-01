#!/usr/bin/env python3
"""
Tahap 4 - Sintesis Multidimensi dan Trade-off (REVISI LENGKAP)
Paket lengkap dengan matriks dominasi, CSV detail, workbook 15 sheet, DOCX bertabel,
SHA-256 verification, dan unit tests.
"""

import sys
import platform
import hashlib
import re
import subprocess
import py_compile
from io import BytesIO
from pathlib import Path
from datetime import datetime
import numpy as np
import pandas as pd
from pandas.testing import assert_frame_equal
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import load_workbook

# ============================================================
# CONFIGURATION
# ============================================================
BASE = Path(__file__).resolve().parents[2]
DATA_ROOT = BASE
TAHAP3 = BASE / "hasil" / "analisis_inferensial"
import os as _os
OUTPUT = Path(_os.environ.get("PIPELINE_OUTPUT_DIR", str(BASE / "hasil" / "sintesis_hasil")))
OUTPUT.mkdir(parents=True, exist_ok=True)

MASTER_DATASET = DATA_ROOT / "data" / "pengujian_final" / "master_dataset_160.csv"
OMNIBUS_CSV = TAHAP3 / "hasil_omnibus_tahap3.csv"
POSTHOC_CSV = TAHAP3 / "hasil_posthoc_tahap3.csv"
CONSISTENCY_OMNIBUS_CSV = TAHAP3 / "hasil_konsistensi_finalerror_omnibus.csv"
TOLERANCE_OMNIBUS_CSV = TAHAP3 / "hasil_proporsi_within_tolerance_omnibus.csv"
TOLERANCE_POSTHOC_CSV = TAHAP3 / "hasil_proporsi_within_tolerance_posthoc.csv"
BRIDGING_CSV = TAHAP3 / "hasil_bridging_deskriptif.csv"

TAHAP3_FILES = [
    OMNIBUS_CSV, POSTHOC_CSV, CONSISTENCY_OMNIBUS_CSV,
    TOLERANCE_OMNIBUS_CSV, TOLERANCE_POSTHOC_CSV, BRIDGING_CSV
]

PARETO_TOL = 1e-9
SCENARIOS = ["Manual Cepat", "Manual Presisi", "Fixed PID", "GS PID"]
SETPOINTS = [15, 20, 25, 30]

VERSION_INFO = {
    "Python": sys.version.split()[0],
    "pandas": pd.__version__,
    "numpy": np.__version__,
    "OS": platform.platform(),
    "Timestamp": datetime.now().isoformat(),
    "Revision": "Paket Lengkap dengan Matriks Dominasi dan Workbook Extended"
}

# ============================================================
# SHA-256 HASHING
# ============================================================

def compute_sha256(filepath):
    sha256_hash = hashlib.sha256()
    if not filepath.exists(): return None
    with open(filepath, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()

def verify_tahap3_integrity(hashes_before):
    mismatches = []
    for filepath, hash_before in hashes_before.items():
        hash_after = compute_sha256(filepath)
        if hash_before != hash_after:
            mismatches.append(filepath.name)
    return len(mismatches) == 0, mismatches

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def dominates_4d(a_vals, b_vals, tol=PARETO_TOL):
    a = np.array(a_vals, dtype=float)
    b = np.array(b_vals, dtype=float)
    not_worse = np.all(a <= b + tol)
    strictly_better = np.any(a < b - tol)
    return not_worse and strictly_better

def test_dominates_4d():
    tests = []
    assert dominates_4d([1, 1, 1, 1], [2, 2, 2, 2]), "Test 1 failed: clear dominance"
    tests.append("Clear dominance")
    assert dominates_4d([1, 2, 2, 2], [2, 2, 2, 2]), "Test 2 failed: partial dominance"
    tests.append("Partial dominance")
    assert not dominates_4d([1, 3, 2, 2], [2, 2, 2, 2]), "Test 3 failed: trade-off"
    tests.append("Trade-off detected")
    assert not dominates_4d([2, 2, 2, 2], [2, 2, 2, 2]), "Test 4 failed: equality"
    tests.append("Equality handled")
    assert not dominates_4d([2 + PARETO_TOL/2, 2, 2, 2], [2, 2, 2, 2]), "Test 5 failed: within tolerance"
    tests.append("Tolerance edge case")
    assert not dominates_4d([2 + PARETO_TOL*2, 2, 2, 2], [2, 2, 2, 2]), "Test 6 failed: outside tolerance"
    tests.append("Outside tolerance")
    return tests

def generate_tradeoff_note(row_a, row_b, outcome_cols):
    a_better = []
    b_better = []
    labels = {
        "MAE_pct": "MAE%",
        "MeanOvershoot_pct": "mean overshoot",
        "MeanDuration_s": "durasi",
        "SD_FinalError_g": "SD FinalError_g"
    }
    for col in outcome_cols:
        if row_a[col] < row_b[col] - PARETO_TOL:
            a_better.append(labels.get(col, col))
        elif row_b[col] < row_a[col] - PARETO_TOL:
            b_better.append(labels.get(col, col))
    
    if not a_better and not b_better:
        return "Nilai keempat outcome setara dalam toleransi numerik."
    
    parts = []
    if a_better:
        parts.append(f"{row_a['Scenario']} lebih rendah pada {', '.join(a_better)}")
    if b_better:
        parts.append(f"{row_b['Scenario']} lebih rendah pada {', '.join(b_better)}")
    
    note = "; ".join(parts) + "."
    if a_better and b_better:
        note += " Tidak terdapat dominasi empat dimensi."
    
    return note

def generate_dominance_matrix(df_summary, setpoint):
    subset = df_summary[df_summary.Setpoint_g == setpoint].copy().sort_values("Scenario").reset_index(drop=True)
    outcomes = ["MAE_pct", "MeanOvershoot_pct", "MeanDuration_s", "SD_FinalError_g"]
    results = []
    for i, row_a in subset.iterrows():
        for j, row_b in subset.iterrows():
            if i >= j: continue
            a_vals = [row_a[col] for col in outcomes]
            b_vals = [row_b[col] for col in outcomes]
            a_dom_b = dominates_4d(a_vals, b_vals)
            b_dom_a = dominates_4d(b_vals, a_vals)
            
            if a_dom_b:
                relation = f"{row_a['Scenario']} dominates {row_b['Scenario']}"
            elif b_dom_a:
                relation = f"{row_b['Scenario']} dominates {row_a['Scenario']}"
            else:
                relation = "trade-off"
            
            results.append({
                "Setpoint_g": setpoint,
                "Scenario_A": row_a["Scenario"],
                "Scenario_B": row_b["Scenario"],
                "A_MAE_pct": row_a["MAE_pct"],
                "A_MeanOvershoot_pct": row_a["MeanOvershoot_pct"],
                "A_MeanDuration_s": row_a["MeanDuration_s"],
                "A_SD_FinalError_g": row_a["SD_FinalError_g"],
                "B_MAE_pct": row_b["MAE_pct"],
                "B_MeanOvershoot_pct": row_b["MeanOvershoot_pct"],
                "B_MeanDuration_s": row_b["MeanDuration_s"],
                "B_SD_FinalError_g": row_b["SD_FinalError_g"],
                "A_dominates_B": a_dom_b,
                "B_dominates_A": b_dom_a,
                "Relation": relation,
                "Tradeoff_note": generate_tradeoff_note(row_a, row_b, outcomes)
            })
    return pd.DataFrame(results)

def calculate_pareto_per_setpoint(df_summary, setpoint):
    subset = df_summary[df_summary.Setpoint_g == setpoint].copy().sort_values("Scenario").reset_index(drop=True)
    outcomes = ["MAE_pct", "MeanOvershoot_pct", "MeanDuration_s", "SD_FinalError_g"]
    results = []
    for i, row_a in subset.iterrows():
        dominated_by = []
        dominates = []
        for j, row_b in subset.iterrows():
            if i == j: continue
            a_vals = [row_a[col] for col in outcomes]
            b_vals = [row_b[col] for col in outcomes]
            if dominates_4d(b_vals, a_vals): dominated_by.append(row_b["Scenario"])
            if dominates_4d(a_vals, b_vals): dominates.append(row_b["Scenario"])
        
        tradeoff_parts = []
        for other in SCENARIOS:
            if other == row_a["Scenario"]: continue
            row_b = subset[subset.Scenario == other].iloc[0]
            if other not in dominated_by and other not in dominates:
                note = generate_tradeoff_note(row_a, row_b, outcomes)
                if "Tidak terdapat dominasi" in note:
                    tradeoff_parts.append(f"vs {other}: {note}")
        
        results.append({
            "Setpoint_g": setpoint,
            "Scenario": row_a["Scenario"],
            "MAE_pct": row_a["MAE_pct"],
            "MeanOvershoot_pct": row_a["MeanOvershoot_pct"],
            "MeanDuration_s": row_a["MeanDuration_s"],
            "SD_FinalError_g": row_a["SD_FinalError_g"],
            "Dominated": len(dominated_by) > 0,
            "DominatedBy": "; ".join(dominated_by) if dominated_by else "—",
            "Dominates": "; ".join(dominates) if dominates else "—",
            "Tradeoff_note": "; ".join(tradeoff_parts) if tradeoff_parts else "—"
        })
    return pd.DataFrame(results)

def rank_within_group(series, ascending=True):
    return series.rank(method="average", ascending=ascending)

def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def num(x, decimals=4):
    if pd.isna(x) or not np.isfinite(x): return "—"
    return f"{float(x):.{decimals}f}".replace(".", ",")


def format_value(value, decimals=4):
    if pd.isna(value):
        return "—"
    if isinstance(value, (bool, np.bool_)):
        return "Ya" if value else "Tidak"
    if isinstance(value, (int, np.integer)):
        return str(int(value))
    if isinstance(value, (float, np.floating)):
        return f"{float(value):.{decimals}f}".replace(".", ",")
    return str(value)


def format_p(value):
    if pd.isna(value):
        return "—"
    return "p < 0,001" if float(value) < 0.001 else f"p = {float(value):.4f}".replace(".", ",")


def add_df_table(doc, caption, dataframe, columns, labels=None, decimals=4):
    labels = labels or {column: column for column in columns}
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(caption)
    run.bold = True
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = labels.get(column, column)
    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            cells[index].text = format_value(row[column], decimals)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for text_paragraph in cell.paragraphs:
                for text_run in text_paragraph.runs:
                    text_run.font.name = "Times New Roman"
                    text_run.font.size = Pt(7.5)
                    text_run.bold = row_index == 0
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    return table


def document_content(path):
    parsed = Document(BytesIO(path.read_bytes()))
    paragraphs = [(paragraph.style.name, paragraph.text.strip()) for paragraph in parsed.paragraphs if paragraph.text.strip()]
    tables = [
        [[cell.text.strip() for cell in row.cells] for row in table.rows]
        for table in parsed.tables
    ]
    return parsed, paragraphs, tables


def correction_method(posthoc_test):
    mapping = {"Dunn-Holm": "Holm", "Tukey HSD": "Tukey HSD", "Games-Howell": "Games-Howell"}
    assert posthoc_test in mapping, f"PostHoc_test tidak dikenal: {posthoc_test}"
    return mapping[posthoc_test]


def assert_close(actual, expected, label):
    assert np.isclose(float(actual), float(expected), rtol=1e-12, atol=1e-12), label

# ============================================================
# MAIN PIPELINE
# ============================================================

print("=" * 70)
print("TAHAP 4 REVISI - SINTESIS MULTIDIMENSI DAN TRADE-OFF")
print("=" * 70)
print()

print(">>> 0. Verifying script compilation and computing SHA-256 hashes")
py_compile.compile(str(Path(__file__).resolve()), doraise=True)
tahap3_hashes_before = {}
for filepath in TAHAP3_FILES:
    if filepath.exists():
        tahap3_hashes_before[filepath] = compute_sha256(filepath)

print(">>> 1. Running unit tests for Pareto dominance")
test_results = test_dominates_4d()
for test in test_results: print(f"  [PASS] {test}")

print(">>> 2. Loading data sources")
df_master = pd.read_csv(MASTER_DATASET)
df_omnibus = pd.read_csv(OMNIBUS_CSV)
df_posthoc = pd.read_csv(POSTHOC_CSV)
df_cons_omni = pd.read_csv(CONSISTENCY_OMNIBUS_CSV)
df_tol_omni = pd.read_csv(TOLERANCE_OMNIBUS_CSV)
df_tol_ph = pd.read_csv(TOLERANCE_POSTHOC_CSV)
df_bridging = pd.read_csv(BRIDGING_CSV)

print(">>> 3. Validating structure and integrity")
assert len(df_master) == 160, f"Expected 160 trials, got {len(df_master)}"
valid_values = df_master.Valid.astype(str).str.upper()
assert valid_values.isin(["TRUE", "1"]).all(), "Not all trials Valid==TRUE"
assert (df_master.StopReason == "TARGET").all(), "Not all trials StopReason==TARGET"

design_check = df_master.groupby(["Scenario", "Setpoint_g"]).size()
assert len(design_check) == 16 and (design_check == 10).all(), "Design not balanced 4x4x10"

assert np.allclose(
    df_master["AbsError_pct"].to_numpy(float),
    df_master["FinalError_pct"].abs().to_numpy(float),
    rtol=1e-12,
    atol=1e-12,
), "AbsError_pct != abs(FinalError_pct)"
df_master["WithinTolerance"] = (df_master.FinalError_pct.abs() <= 5).astype(int)
df_master["HasSettling"] = df_master.SettlingTime_s.notna().astype(int)
assert (df_master.WithinTolerance == df_master.HasSettling).all(), "WithinTolerance <-> SettlingTime availability mismatch"

print(">>> 4. Computing primary summaries per scenario-setpoint")
summary_rows = []
for scenario in SCENARIOS:
    for sp in SETPOINTS:
        subset = df_master[(df_master.Scenario == scenario) & (df_master.Setpoint_g == sp)]
        
        summary_rows.append({
            "Scenario": scenario,
            "Setpoint_g": sp,
            "n": len(subset),
            "MAE_pct": subset.AbsError_pct.mean(),
            "Median_MAE_pct": subset.AbsError_pct.median(),
            "SD_MAE_pct": subset.AbsError_pct.std(ddof=1),
            "Var_MAE_pct": subset.AbsError_pct.var(ddof=1),
            "Q1_MAE_pct": subset.AbsError_pct.quantile(0.25),
            "Q3_MAE_pct": subset.AbsError_pct.quantile(0.75),
            "IQR_MAE_pct": subset.AbsError_pct.quantile(0.75) - subset.AbsError_pct.quantile(0.25),
            
            "MeanOvershoot_pct": subset.MaxOvershoot_pct.mean(),
            "Median_Overshoot_pct": subset.MaxOvershoot_pct.median(),
            "SD_Overshoot_pct": subset.MaxOvershoot_pct.std(ddof=1),
            "Var_Overshoot_pct": subset.MaxOvershoot_pct.var(ddof=1),
            "Q1_Overshoot_pct": subset.MaxOvershoot_pct.quantile(0.25),
            "Q3_Overshoot_pct": subset.MaxOvershoot_pct.quantile(0.75),
            "IQR_Overshoot_pct": subset.MaxOvershoot_pct.quantile(0.75) - subset.MaxOvershoot_pct.quantile(0.25),
            
            "MeanDuration_s": subset.Duration_s.mean(),
            "Median_Duration_s": subset.Duration_s.median(),
            "SD_Duration_s": subset.Duration_s.std(ddof=1),
            "Var_Duration_s": subset.Duration_s.var(ddof=1),
            "Q1_Duration_s": subset.Duration_s.quantile(0.25),
            "Q3_Duration_s": subset.Duration_s.quantile(0.75),
            "IQR_Duration_s": subset.Duration_s.quantile(0.75) - subset.Duration_s.quantile(0.25),
            
            "MeanFinalError_g": subset.FinalError_g.mean(),
            "Median_FinalError_g": subset.FinalError_g.median(),
            "SD_FinalError_g": subset.FinalError_g.std(ddof=1),
            "Var_FinalError_g": subset.FinalError_g.var(ddof=1),
            "Q1_FinalError_g": subset.FinalError_g.quantile(0.25),
            "Q3_FinalError_g": subset.FinalError_g.quantile(0.75),
            "IQR_FinalError_g": subset.FinalError_g.quantile(0.75) - subset.FinalError_g.quantile(0.25),
        })

df_summary = pd.DataFrame(summary_rows)

print(">>> 5. Computing rankings per metric-setpoint")
ranking_rows = []
for sp in SETPOINTS:
    subset = df_summary[df_summary.Setpoint_g == sp].copy()
    subset["Rank_MAE"] = rank_within_group(subset.MAE_pct)
    subset["Rank_Overshoot"] = rank_within_group(subset.MeanOvershoot_pct)
    subset["Rank_Duration"] = rank_within_group(subset.MeanDuration_s)
    subset["Rank_SD"] = rank_within_group(subset.SD_FinalError_g)
    
    for _, row in subset.iterrows():
        ranking_rows.append({
            "Setpoint_g": sp, "Scenario": row.Scenario, "n": row.n,
            "MAE_pct": row.MAE_pct, "Median_MAE_pct": row.Median_MAE_pct, "SD_MAE_pct": row.SD_MAE_pct, 
            "Var_MAE_pct": row.Var_MAE_pct, "Q1_MAE_pct": row.Q1_MAE_pct, "Q3_MAE_pct": row.Q3_MAE_pct, 
            "IQR_MAE_pct": row.IQR_MAE_pct, "Rank_MAE": row.Rank_MAE,
            
            "MeanOvershoot_pct": row.MeanOvershoot_pct, "Median_Overshoot_pct": row.Median_Overshoot_pct, 
            "SD_Overshoot_pct": row.SD_Overshoot_pct, "Var_Overshoot_pct": row.Var_Overshoot_pct, 
            "Q1_Overshoot_pct": row.Q1_Overshoot_pct, "Q3_Overshoot_pct": row.Q3_Overshoot_pct, 
            "IQR_Overshoot_pct": row.IQR_Overshoot_pct, "Rank_Overshoot": row.Rank_Overshoot,
            
            "MeanDuration_s": row.MeanDuration_s, "Median_Duration_s": row.Median_Duration_s, 
            "SD_Duration_s": row.SD_Duration_s, "Var_Duration_s": row.Var_Duration_s, 
            "Q1_Duration_s": row.Q1_Duration_s, "Q3_Duration_s": row.Q3_Duration_s, 
            "IQR_Duration_s": row.IQR_Duration_s, "Rank_Duration": row.Rank_Duration,
            
            "MeanFinalError_g": row.MeanFinalError_g, "Median_FinalError_g": row.Median_FinalError_g,
            "SD_FinalError_g": row.SD_FinalError_g, "Var_FinalError_g": row.Var_FinalError_g, 
            "Q1_FinalError_g": row.Q1_FinalError_g, "Q3_FinalError_g": row.Q3_FinalError_g, 
            "IQR_FinalError_g": row.IQR_FinalError_g, "Rank_SD": row.Rank_SD
        })

df_ranking = pd.DataFrame(ranking_rows)

print(">>> 6. Computing 4D Pareto and Dominance Matrix")
pareto_dfs = []
matrix_dfs = []
for sp in SETPOINTS:
    pareto_dfs.append(calculate_pareto_per_setpoint(df_summary, sp))
    matrix_dfs.append(generate_dominance_matrix(df_summary, sp))

df_pareto = pd.concat(pareto_dfs, ignore_index=True)
df_matrix = pd.concat(matrix_dfs, ignore_index=True)

print(">>> 7. Computing additional and conditional profiles")
# RiseTime_10_90_s
risetime_rows = []
for scenario in SCENARIOS:
    for sp in SETPOINTS:
        subset = df_master[(df_master.Scenario == scenario) & (df_master.Setpoint_g == sp)]
        risetime_rows.append({
            "Scenario": scenario, "Setpoint_g": sp, "RiseTime_role": "Tambahan",
            "RiseTime_n": len(subset),
            "RiseTime_mean": subset.RiseTime_10_90_s.mean(),
            "RiseTime_median": subset.RiseTime_10_90_s.median(),
            "RiseTime_SD": subset.RiseTime_10_90_s.std(ddof=1),
            "RiseTime_Q1": subset.RiseTime_10_90_s.quantile(0.25),
            "RiseTime_Q3": subset.RiseTime_10_90_s.quantile(0.75),
            "RiseTime_IQR": subset.RiseTime_10_90_s.quantile(0.75) - subset.RiseTime_10_90_s.quantile(0.25),
        })
df_risetime = pd.DataFrame(risetime_rows)

# WithinTolerance and SettlingTime_s
tolerance_rows = []
for scenario in SCENARIOS:
    for sp in SETPOINTS:
        subset = df_master[(df_master.Scenario == scenario) & (df_master.Setpoint_g == sp)]
        n_within = subset.WithinTolerance.sum()
        settling_subset = subset[subset.WithinTolerance == 1]
        
        tolerance_rows.append({
            "Scenario": scenario, "Setpoint_g": sp, "Within_role": "Tambahan",
            "Within_total_n": len(subset), "Within_n": n_within, "Within_prop": n_within / len(subset),
            "Settling_role": "Kondisional", "Settling_subset_n": len(settling_subset), 
            "Settling_subset_prop": len(settling_subset) / len(subset),
            "Settling_median": settling_subset.SettlingTime_s.median() if len(settling_subset) > 0 else np.nan,
            "Settling_Q1": settling_subset.SettlingTime_s.quantile(0.25) if len(settling_subset) > 0 else np.nan,
            "Settling_Q3": settling_subset.SettlingTime_s.quantile(0.75) if len(settling_subset) > 0 else np.nan,
            "Settling_IQR": (settling_subset.SettlingTime_s.quantile(0.75) - settling_subset.SettlingTime_s.quantile(0.25)) if len(settling_subset) > 0 else np.nan,
        })
df_tolerance = pd.DataFrame(tolerance_rows)

# Bridging
df_bridging_renamed = df_bridging.rename(columns={
    "n": "Bridging_n", "Total_events": "Bridging_total", "Median": "Bridging_median",
    "IQR_lo": "Bridging_Q1", "IQR_hi": "Bridging_Q3", "Min": "Bridging_min", "Max": "Bridging_max",
    "Prop_nonzero": "Bridging_prop_nonzero"
})
df_bridging_renamed["Bridging_IQR"] = df_bridging_renamed["Bridging_Q3"] - df_bridging_renamed["Bridging_Q1"]
df_bridging_renamed["Bridging_role"] = "Pendukung"

df_additional = df_risetime.merge(df_tolerance, on=["Scenario", "Setpoint_g"])
df_additional = df_additional.merge(df_bridging_renamed, on=["Scenario", "Setpoint_g"], how="left")

# Anotasi inferensial Tahap 3 melalui join berdasarkan metrik dan setpoint.
rise_source = df_omnibus[df_omnibus.Metric == "RiseTime_10_90_s"][
    ["Metric", "Setpoint_g", "Test", "p_raw", "p_holm", "Significant_holm"]
].rename(columns={
    "Test": "RiseTime_Test",
    "p_raw": "RiseTime_p_raw",
    "p_holm": "RiseTime_p_holm",
    "Significant_holm": "RiseTime_Significant_holm",
})
rise_source["RiseTime_source_file"] = OMNIBUS_CSV.name
df_additional = df_additional.merge(
    rise_source.drop(columns="Metric"), on="Setpoint_g", validate="many_to_one"
)

within_source = df_tol_omni[
    ["Setpoint_g", "Test", "p_MonteCarlo", "p_holm", "Significant_holm"]
].rename(columns={
    "Test": "Within_Test",
    "p_MonteCarlo": "Within_p_MonteCarlo",
    "p_holm": "Within_p_holm",
    "Significant_holm": "Within_Significant_holm",
})
within_source["Within_source_file"] = TOLERANCE_OMNIBUS_CSV.name
df_additional = df_additional.merge(within_source, on="Setpoint_g", validate="many_to_one")

expected_rise_tests = {15: "Kruskal-Wallis", 20: "Kruskal-Wallis", 25: "Welch ANOVA", 30: "Welch ANOVA"}
assert df_additional.groupby("Setpoint_g").RiseTime_Test.first().to_dict() == expected_rise_tests

# Keterlacakan p-value dan signifikansi memakai kunci baris sumber.
for _, row in df_additional.iterrows():
    source = df_omnibus[
        (df_omnibus.Metric == "RiseTime_10_90_s")
        & (df_omnibus.Setpoint_g == row.Setpoint_g)
        & (df_omnibus.Test == row.RiseTime_Test)
    ]
    assert len(source) == 1
    source = source.iloc[0]
    assert_close(row.RiseTime_p_raw, source.p_raw, "RiseTime p_raw tidak terlacak")
    assert_close(row.RiseTime_p_holm, source.p_holm, "RiseTime p_holm tidak terlacak")
    assert bool(row.RiseTime_Significant_holm) == bool(source.Significant_holm)

    source = df_tol_omni[df_tol_omni.Setpoint_g == row.Setpoint_g]
    assert len(source) == 1
    source = source.iloc[0]
    assert_close(row.Within_p_MonteCarlo, source.p_MonteCarlo, "Within p_MonteCarlo tidak terlacak")
    assert_close(row.Within_p_holm, source.p_holm, "Within p_holm tidak terlacak")
    assert bool(row.Within_Significant_holm) == bool(source.Significant_holm)

print(">>> 8. Exporting CSV files")
csv_primer = OUTPUT / "tahap4_profil_primer.csv"
csv_pareto = OUTPUT / "tahap4_pareto_per_setpoint.csv"
csv_matrix = OUTPUT / "tahap4_matriks_dominasi.csv"
csv_additional = OUTPUT / "tahap4_profil_tambahan_kondisional.csv"
xlsx_path = OUTPUT / "hasil_lengkap_tahap4.xlsx"
docx_path = OUTPUT / "laporan_tahap4_sintesis_multidimensi.docx"
wt_path = OUTPUT / "walkthrough_tahap4.md"

df_ranking.to_csv(csv_primer, index=False)
df_pareto.to_csv(csv_pareto, index=False)
df_matrix.to_csv(csv_matrix, index=False)
df_additional.to_csv(csv_additional, index=False)

# Metadata effect size mempertahankan struktur p-value sumber.
df_es_omni = df_omnibus[[
    "Metric", "Setpoint_g", "Test", "Statistic", "df1", "df2", "p_raw",
    "EffectSize_name", "EffectSize_value", "p_holm", "Significant_holm"
]].copy()
df_es_omni["correction_method"] = "Holm"
df_es_omni["hypothesis_family"] = "Omnibus"
df_es_omni["source_file"] = OMNIBUS_CSV.name

df_es_ph = df_posthoc[[
    "Metric", "Setpoint_g", "Group_A", "Group_B", "Omnibus_test", "PostHoc_test",
    "p_adjusted", "Significant", "EffectSize_name", "EffectSize_value", "CI_lo", "CI_hi", "Direction"
]].copy()
df_es_ph["correction_method"] = df_es_ph.PostHoc_test.map(correction_method)
df_es_ph["hypothesis_family"] = "Post-hoc"
df_es_ph["source_file"] = POSTHOC_CSV.name

for _, row in df_es_omni.iterrows():
    source = df_omnibus[
        (df_omnibus.Metric == row.Metric)
        & (df_omnibus.Setpoint_g == row.Setpoint_g)
        & (df_omnibus.Test == row.Test)
    ]
    assert len(source) == 1
    source = source.iloc[0]
    assert_close(row.p_raw, source.p_raw, "Omnibus p_raw tidak terlacak")
    assert_close(row.p_holm, source.p_holm, "Omnibus p_holm tidak terlacak")
    assert bool(row.Significant_holm) == bool(source.Significant_holm)

for _, row in df_es_ph.iterrows():
    source = df_posthoc[
        (df_posthoc.Metric == row.Metric)
        & (df_posthoc.Setpoint_g == row.Setpoint_g)
        & (df_posthoc.Group_A == row.Group_A)
        & (df_posthoc.Group_B == row.Group_B)
        & (df_posthoc.PostHoc_test == row.PostHoc_test)
    ]
    assert len(source) == 1
    source = source.iloc[0]
    assert_close(row.p_adjusted, source.p_adjusted, "Post-hoc p_adjusted tidak terlacak")
    assert bool(row.Significant) == bool(source.Significant)

hashes_after = {path.name: compute_sha256(path) for path in TAHAP3_FILES}
assert {path.name: tahap3_hashes_before[path] for path in TAHAP3_FILES} == hashes_after
hash_audit = pd.DataFrame([
    {
        "audit_type": "SHA-256 input Tahap 3",
        "check": filename,
        "status": "PASS",
        "source_file": filename,
        "sha256_before": tahap3_hashes_before[next(path for path in TAHAP3_FILES if path.name == filename)],
        "sha256_after": hashes_after[filename],
        "sha256_match": True,
        "integrity_status": "UNCHANGED",
        "execution_time": VERSION_INFO["Timestamp"],
    }
    for filename in hashes_after
])
base_audit = pd.DataFrame([
    {"audit_type": "Validasi pipeline", "check": check, "status": status,
     "source_file": "—", "sha256_before": "—", "sha256_after": "—",
     "sha256_match": "—", "integrity_status": "—", "execution_time": VERSION_INFO["Timestamp"]}
    for check, status in [
        ("python -m py_compile / kompilasi skrip", "PASS"),
        ("Desain 4 x 4 x 10", "PASS"),
        ("Seluruh Valid = TRUE", "PASS"),
        ("Seluruh StopReason = TARGET", "PASS"),
        ("AbsError_pct = abs(FinalError_pct)", "PASS"),
        ("WithinTolerance = abs(FinalError_pct) <= 5", "PASS"),
        ("Unit test dominasi enam kasus", "PASS"),
        ("Keterlacakan p-value berdasarkan key baris", "PASS"),
        ("Metadata post-hoc berdasarkan PostHoc_test", "PASS"),
        ("Skema dan jumlah baris empat CSV", "PENDING_ARTIFACT_CHECK"),
        ("Workbook 15 sheet", "PENDING_ARTIFACT_CHECK"),
        ("DOCX 16 bagian dan tabel fungsional", "PENDING_ARTIFACT_CHECK"),
        ("Tidak ada grafik Tahap 4", "PENDING_ARTIFACT_CHECK"),
        ("Seluruh keluaran berukuran nonnol", "PENDING_ARTIFACT_CHECK"),
        ("Determinisme semantik dua eksekusi", "PENDING_TWO_RUN_CHECK"),
    ]
])
audit_df = pd.concat([base_audit, hash_audit], ignore_index=True)

print(">>> 9. Creating Excel workbook with 15 sheets")
cols_ws = [
    "Scenario", "Setpoint_g", "Within_role", "Within_total_n", "Within_n", "Within_prop",
    "Within_Test", "Within_p_MonteCarlo", "Within_p_holm", "Within_Significant_holm",
    "Within_source_file", "Settling_role", "Settling_subset_n", "Settling_subset_prop",
    "Settling_median", "Settling_Q1", "Settling_Q3", "Settling_IQR"
]
cols_br = [
    "Scenario", "Setpoint_g", "Bridging_role", "Bridging_n", "Bridging_total",
    "Bridging_median", "Bridging_Q1", "Bridging_Q3", "Bridging_IQR",
    "Bridging_min", "Bridging_max", "Bridging_prop_nonzero"
]
notes = pd.DataFrame([
    ["Ranking", "Dihitung per metrik dan setpoint; tidak dijumlahkan."],
    ["Pareto", "Sintesis deskriptif empat dimensi, bukan uji inferensial."],
    ["Non-dominated", "Bukan sinonim terbaik secara mutlak."],
    ["SettlingTime_s", "Deskriptif kondisional: n tersedia, median, dan IQR; tidak masuk ranking/Pareto."],
    ["BridgingCount", "Pendukung deskriptif, bukan pengamatan visual langsung."],
    ["Grafik", "Grafik final merupakan ruang lingkup Tahap 5."],
], columns=["Topik", "Catatan"])
version_info = pd.DataFrame([{**VERSION_INFO, "Workbook_sheet_count": 15}])

with pd.ExcelWriter(xlsx_path, engine="openpyxl") as writer:
    version_info.to_excel(writer, sheet_name="VersionInfo", index=False)
    audit_df.to_excel(writer, sheet_name="Audit", index=False)
    df_summary.to_excel(writer, sheet_name="PrimarySummaries", index=False)
    df_ranking.to_excel(writer, sheet_name="Rankings", index=False)
    df_pareto.to_excel(writer, sheet_name="Pareto4D", index=False)
    df_matrix.to_excel(writer, sheet_name="Matriks_Dominasi", index=False)
    df_additional.to_excel(writer, sheet_name="AdditionalConditional", index=False)
    df_additional[cols_ws].to_excel(writer, sheet_name="Within_Settling", index=False)
    df_additional[cols_br].to_excel(writer, sheet_name="Bridging", index=False)
    df_es_omni.to_excel(writer, sheet_name="EffectSize_Tahap3_Omnibus", index=False)
    df_es_ph.to_excel(writer, sheet_name="EffectSize_Tahap3_PostHoc", index=False)
    df_omnibus[df_omnibus.Metric.isin(["AbsError_pct", "MaxOvershoot_pct", "Duration_s"])].to_excel(writer, sheet_name="Tahap3_Omnibus_Primer", index=False)
    df_cons_omni.to_excel(writer, sheet_name="Tahap3_Consistency", index=False)
    df_tol_omni.to_excel(writer, sheet_name="Tahap3_Tolerance", index=False)
    notes.to_excel(writer, sheet_name="Catatan_Interpretasi", index=False)

print(">>> 10. Creating full academic DOCX")
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(11)
for style_name in ["Title", "Heading 1", "Heading 2"]:
    doc.styles[style_name].font.name = "Times New Roman"

heading = doc.add_heading("LAPORAN TEKNIS TAHAP 4", 0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph("Sintesis Multidimensi, Trade-off, dan Profil Performa")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
required_headings = []
def add_section(title):
    required_headings.append(title)
    doc.add_heading(title, 1)

add_section("1. Tujuan dan Ruang Lingkup Tahap 4")
doc.add_paragraph("Tahap 4 menyintesis keluaran Tahap 3 menjadi profil performa multidimensi. Tahap ini tidak menghitung ulang uji inferensial, tidak menambah keluarga hipotesis, dan tidak menetapkan juara umum. Grafik formal serta narasi visual tetap menjadi ruang lingkup Tahap 5.")

add_section("2. Desain Eksperimen dan Sumber Data")
doc.add_paragraph("Data terdiri atas 160 trial: empat skenario, empat setpoint massa (15 g, 20 g, 25 g, dan 30 g), dan sepuluh pengulangan tiap kombinasi. Seluruh trial valid dan berhenti karena target. Angka inferensial disalin dari Tahap 3 berdasarkan kunci baris sumber.")

add_section("3. Definisi Empat Skenario")
scenario_def = pd.DataFrame([
    ["Manual Cepat", "Skenario open-loop dengan posisi servo utama tetap 40°. Istilah ‘Manual’ hanya merupakan nama skenario dan tidak berarti pakan dituangkan secara manual oleh manusia."],
    ["Manual Presisi", "Skenario open-loop dengan posisi servo utama tetap 20°. Istilah ‘Manual’ hanya merupakan nama skenario dan tidak berarti pakan dituangkan secara manual oleh manusia."],
    ["Fixed PID", "Kontrol closed-loop Proportional–Integral–Derivative dengan nilai Kp, Ki, dan Kd tetap."],
    ["GS PID", "Kontrol PID dengan gain scheduling berbasis aturan menurut zona error. Nilai Kp, Ki, dan Kd berubah sesuai aturan yang telah ditentukan; metode ini bukan adaptive control."],
], columns=["Skenario", "Definisi operasional"])
add_df_table(doc, "Tabel 1. Definisi skenario penelitian", scenario_def, list(scenario_def.columns))
doc.add_paragraph("Open-loop berarti aksi tidak dikoreksi terus-menerus oleh umpan balik. Closed-loop berarti aksi kendali menggunakan umpan balik keluaran. Definisi struktur ini tidak menyatakan salah satu sistem lebih baik.")

add_section("4. Hierarki dan Definisi Outcome")
outcome_def = pd.DataFrame([
    ["AbsError_pct / MAE%", "Primer", "Galat absolut per trial; mean-nya menjadi Mean Absolute Error kelompok dalam persen."],
    ["MaxOvershoot_pct", "Primer", "Overshoot maksimum terhadap setpoint dalam persen."],
    ["Duration_s", "Primer", "Durasi trial dalam detik."],
    ["SD FinalError_g", "Primer", "Standard deviation atau simpangan baku sampel galat akhir dalam gram."],
    ["RiseTime_10_90_s", "Tambahan", "Waktu kenaikan dari 10% menuju 90% target."],
    ["WithinTolerance", "Tambahan", "Status abs(FinalError_pct) <= 5%."],
    ["SettlingTime_s", "Kondisional", "Waktu settling hanya pada trial dalam toleransi."],
    ["BridgingCount", "Pendukung", "Indikator stagnasi firmware, bukan observasi visual langsung."],
], columns=["Outcome", "Peran", "Penjelasan"])
add_df_table(doc, "Tabel 2. Hierarki dan definisi outcome", outcome_def, list(outcome_def.columns))
doc.add_paragraph("IQR adalah interquartile range atau rentang antarkuartil. Effect size adalah ukuran besar efek. P-value adalah ukuran kompatibilitas data dengan hipotesis nol berdasarkan asumsi uji.")

add_section("5. Metode Sintesis Deskriptif dan Ranking")
doc.add_paragraph("Mean dan simpangan baku sampel (ddof = 1) dipakai sebagai statistik utama yang telah dikunci. Median dan IQR menjadi pendamping. Ranking dihitung terpisah per metrik dan setpoint dari presisi penuh; tidak ada total rank, bobot, atau skor komposit.")

add_section("6. Profil Outcome Primer")
primary = df_ranking[["Setpoint_g", "Scenario", "n", "MAE_pct", "MeanOvershoot_pct", "MeanDuration_s", "SD_FinalError_g"]]
add_df_table(doc, "Tabel 3. Profil outcome primer", primary, list(primary.columns), {
    "Setpoint_g": "SP (g)", "Scenario": "Skenario", "n": "n", "MAE_pct": "MAE%",
    "MeanOvershoot_pct": "Mean overshoot (%)", "MeanDuration_s": "Mean durasi (s)",
    "SD_FinalError_g": "SD galat akhir (g)"})
doc.add_paragraph("Tabel ini memuat point estimate atau nilai taksiran sampel. Arah angka bersifat deskriptif dan tidak otomatis membuktikan perbedaan populasi.")

add_section("7. Ranking Per Metrik dan Setpoint")
ranks = df_ranking[["Setpoint_g", "Scenario", "Rank_MAE", "Rank_Overshoot", "Rank_Duration", "Rank_SD"]]
add_df_table(doc, "Tabel 4. Ranking terpisah setiap outcome primer", ranks, list(ranks.columns))
doc.add_paragraph("Ranking membantu pembacaan profil. Analisis tidak menetapkan juara umum.")

add_section("8. Metode Pareto Empat Dimensi")
doc.add_paragraph("Pareto membandingkan MAE%, mean overshoot, mean durasi, dan SD galat akhir secara bersamaan; seluruhnya diminimalkan. A mendominasi B jika A tidak lebih buruk pada empat outcome dan lebih rendah pada sedikitnya satu, dengan toleransi 1e-9. Non-dominated berarti tidak ditemukan dominator dalam data ini. Dominated berarti terdapat dominator. Trade-off berarti masing-masing skenario lebih rendah pada dimensi berbeda. Status ini deskriptif, bukan inferensial.")

add_section("9. Nilai Pembentuk Pareto")
pareto_values = df_pareto[["Setpoint_g", "Scenario", "MAE_pct", "MeanOvershoot_pct", "MeanDuration_s", "SD_FinalError_g"]]
add_df_table(doc, "Tabel 5. Empat nilai pembentuk Pareto", pareto_values, list(pareto_values.columns))

add_section("10. Status Pareto Per Setpoint")
pareto_status = df_pareto[["Setpoint_g", "Scenario", "Dominated", "Dominates", "DominatedBy"]]
add_df_table(doc, "Tabel 6. Status Pareto dan relasi dominasi", pareto_status, list(pareto_status.columns))
for sp in SETPOINTS:
    names = df_pareto[(df_pareto.Setpoint_g == sp) & (~df_pareto.Dominated)].Scenario.tolist()
    doc.add_paragraph(f"SP{sp}: skenario non-dominated adalah {', '.join(names)}.", style="List Bullet")

add_section("11. Matriks Dominasi dan Trade-off")
matrix_report = df_matrix[["Setpoint_g", "Scenario_A", "Scenario_B", "Relation", "Tradeoff_note"]]
add_df_table(doc, "Tabel 7. Matriks dominasi 24 pasangan unik total", matrix_report, list(matrix_report.columns))
doc.add_paragraph("Empat skenario menghasilkan enam pasangan unik per setpoint. Empat setpoint menghasilkan 24 pasangan unik total. Catatan trade-off dibentuk mekanis dari empat nilai primer tanpa klaim inferensial.")

add_section("12. Audit Fixed PID Versus GS PID")
pid_rows = []
for sp in SETPOINTS:
    row = df_matrix[(df_matrix.Setpoint_g == sp) & (df_matrix.Scenario_A.isin(["Fixed PID", "GS PID"])) & (df_matrix.Scenario_B.isin(["Fixed PID", "GS PID"]))].iloc[0]
    pid_rows.append([sp, row.Relation, row.Tradeoff_note])
pid_audit = pd.DataFrame(pid_rows, columns=["Setpoint_g", "Relation", "Catatan mekanis"])
add_df_table(doc, "Tabel 8. Audit Fixed PID dan GS PID", pid_audit, list(pid_audit.columns))
doc.add_paragraph("GS PID mendominasi Fixed PID pada SP15 dan SP25, sedangkan SP20 dan SP30 menunjukkan trade-off. Pernyataan dibatasi pada point estimate. Perbandingan kedua kontrol PID terhadap dua skenario baseline open-loop, yaitu Manual Cepat dan Manual Presisi, digunakan untuk menjelaskan trade-off akurasi, overshoot, durasi, dan konsistensi.")

add_section("13. Profil Rise Time dan Anotasi Tahap 3")
rise_report = df_additional[["Setpoint_g", "Scenario", "RiseTime_mean", "RiseTime_median", "RiseTime_SD", "RiseTime_Test", "RiseTime_p_holm", "RiseTime_Significant_holm"]].copy()
rise_report["RiseTime_p_holm"] = rise_report.RiseTime_p_holm.map(format_p)
add_df_table(doc, "Tabel 9. RiseTime_10_90_s dan omnibus Tahap 3", rise_report, list(rise_report.columns))
doc.add_paragraph("Kruskal–Wallis adalah uji nonparametrik beberapa kelompok. Welch ANOVA adalah analysis of variance yang tidak mensyaratkan kesamaan varians. SP15–SP20 memakai Kruskal–Wallis; SP25–SP30 memakai Welch ANOVA. Koreksi Holm menyesuaikan p-value secara bertahap untuk keluarga pengujian.")

add_section("14. WithinTolerance, Settling Time, dan BridgingCount")
within_report = df_additional[["Setpoint_g", "Scenario", "Within_n", "Within_total_n", "Within_prop", "Within_p_holm", "Settling_subset_n", "Settling_subset_prop", "Settling_median", "Settling_IQR"]].copy()
within_report["Within_p_holm"] = within_report.Within_p_holm.map(format_p)
add_df_table(doc, "Tabel 10. WithinTolerance dan SettlingTime_s kondisional", within_report, list(within_report.columns))
bridging_report = df_additional[["Setpoint_g", "Scenario", "Bridging_n", "Bridging_total", "Bridging_median", "Bridging_IQR", "Bridging_prop_nonzero"]]
add_df_table(doc, "Tabel 11. BridgingCount deskriptif", bridging_report, list(bridging_report.columns))
doc.add_paragraph("WithinTolerance menggunakan abs(FinalError_pct) <= 5%. SettlingTime_s hanya tersedia pada trial dalam toleransi dan disajikan sebagai n tersedia, median, serta IQR tanpa uji inferensial. BridgingCount tetap indikator stagnasi firmware tanpa inferensi baru.")

add_section("15. Effect Size Tahap 3 dan Audit Reproduktif")
omni_report = df_es_omni[["Metric", "Setpoint_g", "Test", "EffectSize_name", "EffectSize_value", "p_holm", "Significant_holm"]].copy()
omni_report["p_holm"] = omni_report.p_holm.map(format_p)
add_df_table(doc, "Tabel 12. Effect size omnibus Tahap 3", omni_report, list(omni_report.columns))
sig_ph = df_es_ph[df_es_ph.Significant][["Metric", "Setpoint_g", "Group_A", "Group_B", "PostHoc_test", "EffectSize_name", "EffectSize_value", "p_adjusted"]].copy()
sig_ph["p_adjusted"] = sig_ph.p_adjusted.map(format_p)
add_df_table(doc, "Tabel 13. Effect size post-hoc signifikan Tahap 3", sig_ph, list(sig_ph.columns))
add_df_table(doc, "Tabel 14. Audit SHA-256 input Tahap 3", hash_audit, ["source_file", "sha256_before", "sha256_after", "sha256_match", "integrity_status"])
doc.add_paragraph("Epsilon squared, omega squared, f2_W, Hedges g, Cliff delta, dan Cramer's V memiliki skala berbeda sehingga tidak dibandingkan langsung. Dunn–Holm memakai Holm; Tukey HSD adalah honestly significant difference; Games–Howell mengakomodasi ketidaksamaan varians. Hash cocok menunjukkan file Tahap 3 tidak berubah selama pipeline, bukan bukti permission sistem operasi read-only.")

add_section("16. Batasan Interpretasi dan Kesimpulan")
for text in [
    "Hasil tidak signifikan tidak ditafsirkan sebagai bukti kesamaan.",
    "Arah point estimate merupakan arah deskriptif.",
    "Status non-dominated bukan sinonim terbaik secara mutlak.",
    "Status dominated tidak membuktikan inferioritas populasi.",
    "SettlingTime_s kondisional dan rentan bias seleksi subset.",
    "BridgingCount bukan observasi visual langsung.",
    "Tidak ada skor komposit, bobot, total rank, atau juara umum.",
]:
    doc.add_paragraph(text, style="List Bullet")
doc.add_paragraph("Dalam data ini, GS PID dan Manual Cepat non-dominated pada seluruh setpoint; Fixed PID juga non-dominated pada SP20 dan SP30; Manual Presisi dominated pada seluruh setpoint. Kesimpulan terbatas pada data, skenario, setpoint, dan prosedur penelitian ini serta tidak menyatakan satu skenario terbaik untuk seluruh tujuan.")
docx_buffer = BytesIO()
doc.save(docx_buffer)
docx_path.write_bytes(docx_buffer.getvalue())

print(">>> 11. Running substantive artifact validation")
required_csv = {
    csv_primer: ({"Scenario", "Setpoint_g", "MAE_pct", "Rank_MAE", "Rank_Overshoot", "Rank_Duration", "Rank_SD"}, 16),
    csv_pareto: ({"Scenario", "Setpoint_g", "Dominates", "DominatedBy", "Tradeoff_note"}, 16),
    csv_matrix: ({"Scenario_A", "Scenario_B", "Relation", "A_dominates_B", "B_dominates_A", "Tradeoff_note"}, 24),
    csv_additional: ({"RiseTime_Test", "RiseTime_p_holm", "RiseTime_Significant_holm", "RiseTime_source_file", "Within_p_holm", "Within_Significant_holm", "Within_source_file", "RiseTime_role", "Within_role", "Settling_role", "Bridging_role"}, 16),
}
for path, (columns, row_count) in required_csv.items():
    loaded = pd.read_csv(path)
    assert len(loaded) == row_count
    assert columns.issubset(loaded.columns), f"Kolom hilang pada {path.name}: {columns - set(loaded.columns)}"

expected_sheets = ["VersionInfo", "Audit", "PrimarySummaries", "Rankings", "Pareto4D", "Matriks_Dominasi", "AdditionalConditional", "Within_Settling", "Bridging", "EffectSize_Tahap3_Omnibus", "EffectSize_Tahap3_PostHoc", "Tahap3_Omnibus_Primer", "Tahap3_Consistency", "Tahap3_Tolerance", "Catatan_Interpretasi"]
workbook = load_workbook(xlsx_path, read_only=True, data_only=True)
assert workbook.sheetnames == expected_sheets
workbook.close()

parsed_doc, paragraph_content, table_content = document_content(docx_path)
actual_headings = [paragraph.text.strip() for paragraph in parsed_doc.paragraphs if paragraph.style.name.startswith("Heading")]
assert actual_headings == required_headings
headers = [set(table[0]) for table in table_content if table]
required_signatures = [
    {"SP (g)", "Skenario", "MAE%"},
    {"Setpoint_g", "Scenario", "Rank_MAE"},
    {"Setpoint_g", "Scenario", "SD_FinalError_g"},
    {"Setpoint_g", "Scenario", "Dominates", "DominatedBy"},
    {"Setpoint_g", "Scenario_A", "Scenario_B", "Relation"},
    {"Setpoint_g", "Relation", "Catatan mekanis"},
    {"Setpoint_g", "Scenario", "RiseTime_Test"},
    {"Setpoint_g", "Scenario", "Within_n", "Settling_subset_n"},
    {"Setpoint_g", "Scenario", "Bridging_total"},
    {"Metric", "Setpoint_g", "EffectSize_name"},
    {"source_file", "sha256_before", "sha256_after", "sha256_match"},
]
for signature in required_signatures:
    assert any(signature.issubset(header) for header in headers), f"Tabel DOCX wajib hilang: {signature}"

doc_text = "\n".join(text for _, text in paragraph_content) + "\n" + "\n".join(cell for table in table_content for row in table for cell in row)
lower_doc = doc_text.lower()
for pattern in [
    r"gs pid (adalah|merupakan) juara",
    r"skenario terbaik secara mutlak",
    r"closed-loop terbukti unggul secara umum",
    r"dominated berarti inferior(?:\.|$)",
    r"pengumpanan open-loop manual",
    r"baseline manual",
]:
    assert not re.search(pattern, lower_doc)
assert not re.search(r"p\s*=\s*0[,.]0000", lower_doc)
for phrase in [
    "posisi servo utama tetap 40°",
    "posisi servo utama tetap 20°",
    "istilah ‘manual’ hanya merupakan nama skenario",
    "gain scheduling berbasis aturan menurut zona error",
    "metode ini bukan adaptive control",
    "dua skenario baseline open-loop, yaitu manual cepat dan manual presisi",
]:
    assert phrase in lower_doc, f"Redaksi metodologis wajib hilang: {phrase}"
for term in ["proportional–integral–derivative", "mean absolute error", "simpangan baku", "rentang antarkuartil", "non-dominated", "kruskal–wallis", "welch anova", "koreksi holm"]:
    assert term in lower_doc, f"Istilah belum dijelaskan: {term}"

output_paths = [csv_primer, csv_pareto, csv_matrix, csv_additional, xlsx_path, docx_path]
assert all(path.exists() and path.stat().st_size > 0 for path in output_paths)
graphics = [path for path in OUTPUT.iterdir() if path.suffix.lower() in {".png", ".jpg", ".jpeg", ".svg", ".pdf"}]
assert not graphics, f"Keluaran grafik ditemukan: {graphics}"
assert hashes_after == {path.name: compute_sha256(path) for path in TAHAP3_FILES}

with wt_path.open("w", encoding="utf-8") as handle:
    handle.write("# Walkthrough Tahap 4 — Sintesis Multidimensi\n\n")
    handle.write("## Status Verifikasi\n\n")
    handle.write("- `python -m py_compile` / kompilasi skrip: **PASS**.\n")
    handle.write("- Skema dan jumlah baris empat CSV (16 / 16 / 24 / 16): **PASS**.\n")
    handle.write("- Perhitungan primer, rank, Pareto, dan matriks: **PASS**.\n")
    handle.write("- Keterlacakan p-value berdasarkan key baris: **PASS**.\n")
    handle.write("- Pemeriksaan isi DOCX 16 bagian dan tabel fungsional melalui `python-docx`: **PASS**.\n")
    handle.write("- Struktur workbook 15 sheet: **PASS**.\n")
    handle.write("- Seluruh keluaran berukuran nonnol: **PASS**.\n")
    handle.write("- Tidak ada grafik Tahap 4: **PASS**.\n")
    handle.write("- Determinisme semantik dua eksekusi: **PENDING_TWO_RUN_CHECK**.\n")
    handle.write("- File Tahap 3 tidak berubah: **PASS**.\n\n")
    handle.write("## Struktur Dominasi\n\n- 6 pasangan unik per setpoint.\n- 24 pasangan unik untuk seluruh empat setpoint.\n\n")
    handle.write("## Anotasi Inferensial Tambahan\n\n- RiseTime SP15–SP20: Kruskal–Wallis.\n- RiseTime SP25–SP30: Welch ANOVA.\n- WithinTolerance: uji Monte Carlo exact pendekatan Fisher–Freeman–Halton dari Tahap 3.\n\n")
    handle.write("## Audit SHA-256 Input Tahap 3\n\n| File | SHA-256 sebelum | SHA-256 sesudah | Match | Integrity status |\n|---|---|---|---|---|\n")
    for filename in hashes_after:
        before = tahap3_hashes_before[next(path for path in TAHAP3_FILES if path.name == filename)]
        handle.write(f"| {filename} | `{before}` | `{hashes_after[filename]}` | TRUE | UNCHANGED |\n")
    handle.write("\nHash cocok menunjukkan file tidak berubah selama pipeline; bukan bukti permission filesystem read-only.\n\n")
    handle.write("## Keluaran Resmi\n\n| File | Ukuran (byte) |\n|---|---:|\n")
    for path in output_paths:
        handle.write(f"| {path.name} | {path.stat().st_size} |\n")
    handle.write("\nWorkbook berisi 15 sheet. Tidak ada grafik Tahap 4.\n\n## Audit Fixed PID Versus GS PID\n\n")
    for _, row in pid_audit.iterrows():
        handle.write(f"- **SP{int(row.Setpoint_g)}:** {row.Relation}. {row['Catatan mekanis']}\n")
    handle.write("\n## Batas Interpretasi\n\nRanking dan Pareto merupakan sintesis deskriptif. Analisis tidak menetapkan juara umum. Hasil tidak signifikan tidak ditafsirkan sebagai kesamaan. SettlingTime_s tetap kondisional dan BridgingCount tetap pendukung deskriptif.\n")

assert wt_path.stat().st_size > 0


def workbook_semantics(path):
    with pd.ExcelFile(BytesIO(path.read_bytes())) as book:
        sheet_names = list(book.sheet_names)
        semantic = {}
        for sheet_name in sheet_names:
            frame = pd.read_excel(book, sheet_name=sheet_name, dtype=object)
            excluded = [column for column in frame.columns if str(column).lower() in {"timestamp", "execution_time"}]
            semantic[sheet_name] = frame.drop(columns=excluded)
    return sheet_names, semantic


def assert_workbook_equal(first, second):
    first_names, first_frames = first
    second_names, second_frames = second
    assert first_names == second_names
    for sheet_name in first_names:
        assert first_frames[sheet_name].shape == second_frames[sheet_name].shape
        assert list(first_frames[sheet_name].columns) == list(second_frames[sheet_name].columns)
        assert_frame_equal(first_frames[sheet_name], second_frames[sheet_name], check_dtype=False, check_exact=True)


if "--determinism-child" not in sys.argv:
    csv_hashes_first = {path.name: compute_sha256(path) for path in [csv_primer, csv_pareto, csv_matrix, csv_additional]}
    frames_first = {
        "ranking": pd.read_csv(csv_primer),
        "pareto": pd.read_csv(csv_pareto),
        "matrix": pd.read_csv(csv_matrix),
    }
    workbook_first = workbook_semantics(xlsx_path)
    _, paragraphs_first, tables_first = document_content(docx_path)

    completed = subprocess.run(
        [sys.executable, str(Path(__file__).resolve()), "--determinism-child"],
        cwd=str(OUTPUT),
        capture_output=True,
        text=True,
        check=False,
    )
    assert completed.returncode == 0, f"Eksekusi determinisme kedua gagal:\n{completed.stdout}\n{completed.stderr}"

    csv_hashes_second = {path.name: compute_sha256(path) for path in [csv_primer, csv_pareto, csv_matrix, csv_additional]}
    assert csv_hashes_first == csv_hashes_second
    assert_frame_equal(frames_first["ranking"], pd.read_csv(csv_primer), check_dtype=False, check_exact=True)
    assert_frame_equal(frames_first["pareto"], pd.read_csv(csv_pareto), check_dtype=False, check_exact=True)
    assert_frame_equal(frames_first["matrix"], pd.read_csv(csv_matrix), check_dtype=False, check_exact=True)
    assert_workbook_equal(workbook_first, workbook_semantics(xlsx_path))
    _, paragraphs_second, tables_second = document_content(docx_path)
    assert paragraphs_first == paragraphs_second
    assert tables_first == tables_second

    workbook = load_workbook(xlsx_path)
    audit_sheet = workbook["Audit"]
    headers = {cell.value: index + 1 for index, cell in enumerate(audit_sheet[1])}
    for row_index in range(2, audit_sheet.max_row + 1):
        status_cell = audit_sheet.cell(row_index, headers["status"])
        if status_cell.value in {"PENDING_ARTIFACT_CHECK", "PENDING_TWO_RUN_CHECK"}:
            status_cell.value = "PASS"
    workbook.save(xlsx_path)

    walkthrough_text = wt_path.read_text(encoding="utf-8")
    walkthrough_text = walkthrough_text.replace(
        "Determinisme semantik dua eksekusi: **PENDING_TWO_RUN_CHECK**.",
        "Determinisme semantik dua eksekusi: **PASS**.",
    )
    walkthrough_text = re.sub(
        r"\| hasil_lengkap_tahap4\.xlsx \| \d+ \|",
        f"| hasil_lengkap_tahap4.xlsx | {xlsx_path.stat().st_size} |",
        walkthrough_text,
    )
    wt_path.write_text(walkthrough_text, encoding="utf-8")
    assert hashes_after == {path.name: compute_sha256(path) for path in TAHAP3_FILES}
    print("PASS: determinisme dua eksekusi, integritas Tahap 3, dan seluruh artefak tervalidasi.")
else:
    print("PASS: eksekusi child determinisme selesai.")
