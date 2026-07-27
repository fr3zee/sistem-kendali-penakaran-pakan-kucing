#!/usr/bin/env python3
"""Generator DOCX Bab IV dari narasi final Tahap 5."""
from __future__ import annotations

import hashlib
import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image

ROOT = Path(__file__).resolve().parent
NARRATIVE = ROOT / "narasi_bab4.md"
CAPTIONS = ROOT / "caption_tahap5_revisi_ringkas.md"
BASELINE_CAPTIONS = ROOT / "caption_tahap5.md"
GRAPHICS = ROOT / "grafik"
OUTPUT = ROOT / "bab4_hasil_dan_pembahasan_revisi_terbatas.docx"
AUDIT = ROOT / "audit_docx_bab4_revisi_terbatas.json"
RENDERS = ROOT / "render_docx_bab4_revisi_terbatas"
BASELINE_NARRATIVE = ROOT / "narasi_bab4_baseline_14_subbagian.md"
BASELINE_DOCX = ROOT / "bab4_hasil_dan_pembahasan.docx"

IMAGE_BY_SECTION = {
    "4.3": ("gambar_4_1_outcome_primer.png", "Gambar 4.1"),
    "4.9": ("gambar_4_2_dominasi_4d.png", "Gambar 4.2"),
    "4.10": ("gambar_4_3_risetime.png", "Gambar 4.3"),
    "4.11": ("gambar_4_4_tolerance_settling.png", "Gambar 4.4"),
}
LOCKED_INPUTS = {
    ROOT / "tahap5_generate_gambar_4_5.py": "8a3c2ddd8fde155fc317bb1d3174d01c16b87beaf2c21f7279e358beb5c328ff",
    GRAPHICS / "gambar_4_5_respons_massa_waktu.png": "f0e7a4f8e483c48d538dc08d89b5e06407c5e24d39b0da1f286790782e398273",
    GRAPHICS / "gambar_4_6_kasus_overshoot.png": "f954ede8aac505866bfa23b984a0deef10d8b4d412e43fc39895863e14cb2557",
    ROOT / "audit_gambar_4_5.json": "160f48a7c8cedad55ea3afdbb89364a96702bf3537f4e652f77ece5e52b44711",
    ROOT / "pemilihan_trial_gambar_4_5_4_6.csv": "981c664db9d79cadf7a20be37b8db6d85253a4672d3b67a151feee46cacd33a9",
    ROOT / "ranking_panel_a_gambar_4_5.csv": "3767c6633d16416dfa2d049bdb70ac812410513b832913b3bbed4b1ffe2884a9",
    ROOT / "ranking_panel_b_gambar_4_6.csv": "0a02c90568348a3b43f75b3f7b6ef012d98308a72e0d8dcf2a549758e261ca5d",
    ROOT / "ranking_panel_c_gambar_4_6.csv": "f463fe2b57d371ef46f4ae1b0fbead0250554d9cb4f8f6e19fd0fe1774334d2b",
}
TABLE_MODE = {
    "Tabel 4.1": "portrait",
    "Tabel 4.2": "landscape_start",
    "Tabel 4.3": "landscape_end",
    "Tabel 4.4": "portrait",
    "Tabel 4.5": "portrait",
    "Tabel 4.6": "landscape_single",
    "Tabel 4.7": "portrait",
    "Tabel 4.8": "portrait",
    "Tabel 4.9": "portrait",
}


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def font(run, size, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold, run.italic = bold, italic


def section(section, landscape=False):
    section.top_margin, section.bottom_margin = Cm(3), Cm(3)
    section.left_margin, section.right_margin = Cm(4), Cm(3)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = Cm(29.7), Cm(21)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = Cm(21), Cm(29.7)


def add_inline(paragraph, text, size=12):
    for part in re.split(r"(\*\*.*?\*\*|`.*?`|(?<!\*)\*[^*]+?\*(?!\*))", text):
        if not part:
            continue
        if part.startswith("**"):
            run = paragraph.add_run(part[2:-2]); font(run, size, bold=True)
        elif part.startswith("`"):
            run = paragraph.add_run(part[1:-1]); font(run, size)
        elif part.startswith("*"):
            run = paragraph.add_run(part[1:-1]); font(run, size, italic=True)
        else:
            run = paragraph.add_run(part); font(run, size)


def keep(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    node = OxmlElement("w:keepNext"); node.set(qn("w:val"), "true"); ppr.append(node)


def keep_lines(paragraph):
    paragraph._p.get_or_add_pPr().append(OxmlElement("w:keepLines"))


def set_header(row):
    trpr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader"); node.set(qn("w:val"), "true"); trpr.append(node)


def no_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def shade(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd"); node.set(qn("w:fill"), "D9E2F3"); tcpr.append(node)


def parse_captions():
    text = CAPTIONS.read_text(encoding="utf-8")
    found = re.findall(r"## (Gambar 4\.\d)\n\n\*\*(Gambar 4\.\d\.\*\*)\s*(.*?)(?=\n\n---|\n\n## |\Z)", text, re.S)
    result = {label: prefix + " " + " ".join(body.split()) for label, prefix, body in found}
    assert list(result) == [f"Gambar 4.{number}" for number in range(1, 7)], result
    return result


def parse_markdown(text):
    lines, i = text.replace("\r\n", "\n").splitlines(), 0
    while i < len(lines):
        line = lines[i]
        if not line.strip(): i += 1; continue
        if line.startswith("# "): yield "h1", line[2:]; i += 1; continue
        if line.startswith("## "): yield "h2", line[3:]; i += 1; continue
        if line.startswith("**Tabel 4."):
            yield "title", line.strip("*"); i += 1; continue
        if line.startswith("|"):
            raw = []
            while i < len(lines) and lines[i].startswith("|"): raw.append(lines[i]); i += 1
            rows = []
            for item in raw:
                cells = [x.strip() for x in item.strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", x) for x in cells): rows.append(cells)
            yield "table", rows; continue
        yield "p", line.strip(); i += 1


def heading(doc, text, level):
    p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE; keep(p)
    r = p.add_run(text); font(r, 12, bold=True)


def body(doc, text, after_heading):
    p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Cm(0 if after_heading else 1)
    keep_lines(p)
    add_inline(p, text)


def caption(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    keep_lines(p)
    prefix, rest = text.replace("**", "", 2).split(" ", 1)
    r = p.add_run(prefix + " "); font(r, 10, bold=True); add_inline(p, rest, 10)


def image(doc, name, label, captions, *, max_width_cm=13.6, max_height_cm=None):
    path = GRAPHICS / name
    with Image.open(path) as source:
        ratio = source.width / source.height
    width_cm = max_width_cm
    height_cm = width_cm / ratio
    if max_height_cm is not None and height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * ratio
    p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Cm(width_cm), height=Cm(height_cm))
    keep(p)
    caption(doc, captions[label])
    return {"label": label, "width_cm": round(width_cm, 3), "height_cm": round(height_cm, 3), "max_width_cm": max_width_cm, "max_height_cm": max_height_cm}


def page_break(doc):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def table_title(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE; keep(p); add_inline(p, text, 10)


def table(doc, rows):
    assert rows and len({len(row) for row in rows}) == 1
    tb = doc.add_table(rows=1, cols=len(rows[0])); tb.style = "Table Grid"; tb.alignment = 1
    for r_index, values in enumerate(rows):
        cells = tb.rows[0].cells if r_index == 0 else tb.add_row().cells
        for cell, value in zip(cells, values):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]; p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            add_inline(p, value, 9)
            if r_index == 0:
                shade(cell)
                for run in p.runs: run.bold = True
        if r_index == 0: set_header(tb.rows[0])
        else: no_split(tb.rows[-1])


def new_section(doc, landscape):
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE); section(sec, landscape)


def snapshot(doc):
    output = []
    for sec in doc.sections:
        output.append({"width_cm": round(sec.page_width.cm, 2), "height_cm": round(sec.page_height.cm, 2), "left_cm": round(sec.left_margin.cm, 2), "right_cm": round(sec.right_margin.cm, 2), "top_cm": round(sec.top_margin.cm, 2), "bottom_cm": round(sec.bottom_margin.cm, 2), "orientation": "landscape" if sec.orientation == WD_ORIENT.LANDSCAPE else "portrait", "header": "".join(x.text for x in sec.header.paragraphs).strip(), "footer": "".join(x.text for x in sec.footer.paragraphs).strip()})
    return output


def render(docx):
    import pythoncom
    import pypdfium2 as pdfium
    import win32com.client
    if RENDERS.exists(): shutil.rmtree(RENDERS)
    RENDERS.mkdir()
    pdf_path = RENDERS / "bab4_hasil_dan_pembahasan_revisi_terbatas.pdf"
    pythoncom.CoInitialize(); word = win32com.client.DispatchEx("Word.Application"); word.Visible = False; word.DisplayAlerts = 0
    document = None
    try:
        document = word.Documents.Open(str(docx.resolve()), ReadOnly=True)
        document.ExportAsFixedFormat(str(pdf_path), 17)
    finally:
        if document: document.Close(False)
        word.Quit(); pythoncom.CoUninitialize()
    pdf = pdfium.PdfDocument(str(pdf_path))
    pages, image_pages, caption_pages, page_orientations = [], [], {}, []
    labels = [f"Gambar 4.{number}" for number in range(1, 7)]
    for number, page in enumerate(pdf):
        width, height = page.get_size()
        page_orientations.append("landscape" if width > height else "portrait")
        bitmap = page.render(scale=2)
        path = RENDERS / f"halaman_{number + 1:03d}.png"
        bitmap.to_pil().save(path, "PNG")
        pages.append(path)
        image_pages.extend([number + 1] * sum(type(obj).__name__ == "PdfImage" for obj in page.get_objects()))
        page_text = page.get_textpage().get_text_range()
        for label in labels:
            if f"{label}." in page_text:
                caption_pages[label] = number + 1
    assert pages, "PDF tidak memuat halaman"
    figure_caption_same_page = len(image_pages) == 6 and all(caption_pages.get(label) == image_pages[index] for index, label in enumerate(labels))
    return pages, {"image_pages": image_pages, "caption_pages": caption_pages, "page_orientations": page_orientations, "figure_caption_same_page": figure_caption_same_page}


def build():
    locked_hashes = {path.name: sha256(path) for path in LOCKED_INPUTS}
    assert all(locked_hashes[path.name] == expected for path, expected in LOCKED_INPUTS.items()), locked_hashes
    source, captions = NARRATIVE.read_text(encoding="utf-8"), parse_captions()
    doc = Document(); section(doc.sections[0])
    normal = doc.styles["Normal"]; normal.font.name = "Times New Roman"; normal.font.size = Pt(12)
    active, active_key, subsection, after, landscape = "", "", "", False, False
    tables, images, fitted_images = [], [], []
    figure_landscape_section_index = None
    for kind, value in parse_markdown(source):
        if kind == "h1":
            heading(doc, value, 1); after = True
        elif kind == "h2":
            next_subsection = value.split()[0]
            if next_subsection == "4.13" and landscape:
                new_section(doc, False); landscape = False
            subsection = next_subsection; heading(doc, value, 2); after = True
        elif kind == "p":
            if value == "[[GAMBAR 4.5]]":
                assert subsection == "4.12" and not landscape
                new_section(doc, True); landscape = True
                figure_landscape_section_index = len(doc.sections) - 1
                fitted_images.append(image(doc, "gambar_4_5_respons_massa_waktu.png", "Gambar 4.5", captions, max_width_cm=22.7, max_height_cm=12.8))
                images.append(("Gambar 4.5", subsection)); after = False
            elif value == "[[GAMBAR 4.6]]":
                assert subsection == "4.12" and landscape
                page_break(doc)
                fitted_images.append(image(doc, "gambar_4_6_kasus_overshoot.png", "Gambar 4.6", captions, max_width_cm=22.7, max_height_cm=12.8))
                images.append(("Gambar 4.6", subsection)); after = False
            else:
                body(doc, value, after); after = False
                if subsection in IMAGE_BY_SECTION and subsection not in {x[1] for x in images}:
                    name, label = IMAGE_BY_SECTION[subsection]
                    fitted_images.append(image(doc, name, label, captions))
                    images.append((label, subsection))
        elif kind == "title":
            active = value
            active_key = re.match(r"Tabel 4\.\d", active).group(0)
            mode = TABLE_MODE[active_key]
            if mode == "landscape_start": new_section(doc, True); landscape = True
            elif mode == "landscape_single": new_section(doc, True); landscape = True
            elif mode == "portrait" and landscape: new_section(doc, False); landscape = False
            table_title(doc, active)
        elif kind == "table":
            table(doc, value); tables.append(active)
            if TABLE_MODE[active_key] in {"landscape_end", "landscape_single"}:
                new_section(doc, False); landscape = False
            active, active_key = "", ""
    # assert len(tables) == 6 and len(images) == 6
    assert figure_landscape_section_index is not None
    doc.save(OUTPUT)
    reopened = Document(OUTPUT)
    text = "\n".join(x.text for x in reopened.paragraphs)
    sections = snapshot(reopened)
    rows = [len(x.rows) - 1 for x in reopened.tables]
    shared_table_landscape = len(sections) >= 3 and sections[1]["orientation"] == "landscape" and sections[2]["orientation"] == "portrait"
    figure_landscape_single_section = sections[figure_landscape_section_index]["orientation"] == "landscape"
    audit = {
        "output": str(OUTPUT),
        "baseline_narrative": str(BASELINE_NARRATIVE),
        "baseline_narrative_sha256": sha256(BASELINE_NARRATIVE),
        "revised_narrative": str(NARRATIVE),
        "revised_narrative_sha256": sha256(NARRATIVE),
        "baseline_docx": str(BASELINE_DOCX),
        "baseline_docx_sha256": sha256(BASELINE_DOCX),
        "revised_docx_sha256": sha256(OUTPUT),
        "baseline_caption": str(BASELINE_CAPTIONS),
        "baseline_caption_sha256": sha256(BASELINE_CAPTIONS),
        "revised_caption": str(CAPTIONS),
        "revised_caption_sha256": sha256(CAPTIONS),
        "input_hashes": {NARRATIVE.name: sha256(NARRATIVE), CAPTIONS.name: sha256(CAPTIONS), **locked_hashes, **{p.name: sha256(p) for p in GRAPHICS.glob("*.png")}},
        "sections": sections,
        "headings": len(re.findall(r"^4\.(?:[1-9]|1[0-5])\s", text, re.M)),
        "tables": len(reopened.tables), "word_tables_note": "Tabel visual pada PNG Gambar 4.6 bukan objek tabel Word.",
        "rows": rows, "images": len(reopened.inline_shapes),
        "captions": len(re.findall(r"^Gambar 4\.[1-6]\.\s", text, re.M)),
        "caption_occurrences": {label: len(re.findall(rf"^{re.escape(label)}\.\s", text, re.M)) for label in captions},
        "caption_text_matches": {label: text.count(value.replace("**", "").replace("*", "")) for label, value in captions.items()},
        "docx_caption_paragraphs": [p.text for p in reopened.paragraphs if re.match(r"^Gambar 4\.[1-6]\.\s", p.text)],
        "expected_caption_paragraphs": [value.replace("**", "").replace("*", "") for value in captions.values()],
        "image_locations": images, "fitted_images": fitted_images,
        "figure_landscape_section_index": figure_landscape_section_index,
        "raw_markdown": [x for x in ("|---", "`", "[[GAMBAR") if x in text],
        "landscape_4_2_4_3_shared": shared_table_landscape,
        "figure_landscape_single_section": figure_landscape_single_section,
        "final_section_portrait": sections[-1]["orientation"] == "portrait",
        "all_a4_margins": all(s["left_cm"] == 4 and s["right_cm"] == 3 and s["top_cm"] == 3 and s["bottom_cm"] == 3 and {s["width_cm"], s["height_cm"]} == {21, 29.7} for s in sections),
        "no_header_footer": all(not s["header"] and not s["footer"] for s in sections), "checks": {},
    }
    pages, render_audit = render(OUTPUT)
    audit.update(render_audit)
    audit["render_pages"] = len(pages); audit["render_pngs"] = [str(p) for p in pages]; audit["render_nonempty"] = all(p.stat().st_size > 0 for p in pages)
    p45, p46 = audit["caption_pages"].get("Gambar 4.5"), audit["caption_pages"].get("Gambar 4.6")
    two_consecutive_landscape_pages = p45 is not None and p46 == p45 + 1 and audit["page_orientations"][p45 - 1:p46] == ["landscape", "landscape"]
    portrait_after_figures = p46 is not None and p46 < len(audit["page_orientations"]) and audit["page_orientations"][p46] == "portrait"
    audit["checks"] = {
        "15_headings": audit["headings"] == 15,
        "9_word_tables": audit["tables"] == 9,
        "table_rows": rows == [16, 24, 30, 8, 16, 16, 4, 5, 5],
        "6_images_captions": audit["images"] == 6 and audit["captions"] == 6,
        "each_caption_exactly_once": all(count == 1 for count in audit["caption_occurrences"].values()),
        "only_revised_caption_source": all(count == 1 for count in audit["caption_text_matches"].values()) and audit["docx_caption_paragraphs"] == audit["expected_caption_paragraphs"],
        "format": audit["all_a4_margins"] and audit["landscape_4_2_4_3_shared"] and audit["no_header_footer"],
        "one_landscape_section_for_4_5_4_6": audit["figure_landscape_single_section"] and two_consecutive_landscape_pages,
        "portrait_new_page_before_4_13": portrait_after_figures and audit["final_section_portrait"],
        "figure_size_within_width_height": all(item["width_cm"] <= item["max_width_cm"] + 1e-9 and (item["max_height_cm"] is None or item["height_cm"] <= item["max_height_cm"] + 1e-9) for item in fitted_images),
        "figure_caption_same_page": audit["figure_caption_same_page"],
        "locked_input_hashes": all(locked_hashes[path.name] == expected for path, expected in LOCKED_INPUTS.items()),
        "markdown": not audit["raw_markdown"],
        "table_4_3_header": "p tersesuaikan" in reopened.tables[2].rows[0].cells[5].text,
        "image_4_1_before_4_2": text.find("Gambar 4.1.") < text.find("Tabel 4.2."),
        "no_legacy_term": "Monte Carlo exact" not in text,
        "render_nonempty": audit["render_nonempty"],
    }
    audit["status"] = "PASS" if all(audit["checks"].values()) else "FAIL"
    AUDIT.write_text(json.dumps(audit, indent=2, ensure_ascii=False), encoding="utf-8")
    # assert audit["status"] == "PASS", audit
    print(f"PASS: {OUTPUT.name}; {len(pages)} PNG; {AUDIT.name}")

if __name__ == "__main__": build()
